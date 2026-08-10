"""Extract text from uploaded files and split content into searchable chunks."""

import base64
import io
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openai import OpenAI
from PIL import Image

from backend.config import CHUNK_OVERLAP, CHUNK_SIZE, OPENAI_API_KEY, VISION_MODEL
from backend.prompts.vision import VISION_PROMPT

MIN_TEXT_CHARS = 50


def _get_openai_client() -> OpenAI:
    return OpenAI(api_key=OPENAI_API_KEY)


def _vision_extract_text(image_bytes: bytes, mime_type: str = "image/png") -> str:
    client = _get_openai_client()
    encoded = base64.b64encode(image_bytes).decode("utf-8")
    response = client.chat.completions.create(
        model=VISION_MODEL,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": VISION_PROMPT},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime_type};base64,{encoded}"},
                    },
                ],
            }
        ],
        max_tokens=4096,
    )
    return (response.choices[0].message.content or "").strip()


def _pdf_page_to_image_bytes(page: fitz.Page, dpi: int = 200) -> bytes:
    pixmap = page.get_pixmap(dpi=dpi)
    return pixmap.tobytes("png")


def _load_pdf(path: Path, source_name: str) -> tuple[list[Document], int]:
    documents: list[Document] = []
    pdf = fitz.open(path)

    for page_index in range(len(pdf)):
        page = pdf[page_index]
        page_number = page_index + 1
        text = page.get_text("text").strip()

        if len(text) < MIN_TEXT_CHARS:
            try:
                image_bytes = _pdf_page_to_image_bytes(page)
                text = _vision_extract_text(image_bytes, "image/png")
            except Exception:
                text = text or ""

        if text:
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": source_name,
                        "page": page_number,
                        "file_type": "pdf",
                    },
                )
            )

    page_count = len(pdf)
    pdf.close()
    return documents, page_count


def _load_docx(path: Path, source_name: str) -> tuple[list[Document], int]:
    doc = DocxDocument(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    if not text:
        return [], 1
    return [
        Document(
            page_content=text,
            metadata={"source": source_name, "page": 1, "file_type": "docx"},
        )
    ], 1


def _load_text(path: Path, source_name: str, file_type: str) -> tuple[list[Document], int]:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        return [], 1
    return [
        Document(
            page_content=text,
            metadata={"source": source_name, "page": 1, "file_type": file_type},
        )
    ], 1


def _load_image(path: Path, source_name: str) -> tuple[list[Document], int]:
    suffix = path.suffix.lower().lstrip(".")
    mime_map = {
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "png": "image/png",
        "webp": "image/webp",
        "gif": "image/gif",
        "bmp": "image/bmp",
        "tiff": "image/tiff",
    }
    mime_type = mime_map.get(suffix, "image/png")
    image_bytes = path.read_bytes()

    with Image.open(io.BytesIO(image_bytes)) as img:
        if img.mode not in ("RGB", "RGBA"):
            buffer = io.BytesIO()
            img.convert("RGB").save(buffer, format="PNG")
            image_bytes = buffer.getvalue()
            mime_type = "image/png"

    text = _vision_extract_text(image_bytes, mime_type)
    if not text:
        return [], 1
    return [
        Document(
            page_content=text,
            metadata={"source": source_name, "page": 1, "file_type": suffix},
        )
    ], 1


def load_document(path: Path, source_name: str | None = None) -> tuple[list[Document], int]:
    source = source_name or path.name
    suffix = path.suffix.lower().lstrip(".")

    if suffix == "pdf":
        return _load_pdf(path, source)
    if suffix == "docx":
        return _load_docx(path, source)
    if suffix in {"txt", "md"}:
        return _load_text(path, source, suffix)
    if suffix in {"png", "jpg", "jpeg", "webp", "gif", "bmp", "tiff"}:
        return _load_image(path, source)

    raise ValueError(f"Unsupported file type: .{suffix}")


def chunk_documents(documents: list[Document], document_id: int) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
    )
    chunks = splitter.split_documents(documents)
    for index, chunk in enumerate(chunks):
        chunk.metadata["document_id"] = document_id
        chunk.metadata["chunk_index"] = index
    return chunks
